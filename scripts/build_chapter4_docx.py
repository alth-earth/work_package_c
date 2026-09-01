#!/usr/bin/env python3
"""Build Chapter 4 Word document from existing evidence.

Generates ``docs/CHAPTER4.docx`` from the figures, tables and data already
produced by the algorithm-comparison pipeline.  No fabricated numbers; every
figure and table is referenced from the underlying artefacts under
``.runtime/experiments/c-algorithm-comparison-*``.

Run with::

    uv run --with python-docx python scripts/build_chapter4_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree

REPO_ROOT = Path(__file__).resolve().parents[1]
FIGURES = (
    REPO_ROOT.parent / ".runtime" / "experiments" / "c-algorithm-comparison-summary" / "figures"
)
REFERENCE_DOCX = Path("/mnt/c/Users/asd233/Downloads/第一章(1).docx")
OUTPUT = REPO_ROOT / "docs" / "第四章.docx"


def _normalise_labels(text: str) -> str:
    """Match the compact numbering convention used by Chapter 1."""
    text = re.sub(r"第\s+4\s+章", "第4章", text)
    text = re.sub(r"([图表])\s+4-", r"\g<1>4-", text)
    text = re.sub(r"式\s+\(4-", "式(4-", text)
    return text


def _set_font(style, east_asia: str, latin: str, size_pt: float, *, bold: bool = False) -> None:
    style.font.name = latin
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = None
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def _clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _configure_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_font(normal, "宋体", "Times New Roman", 10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = Pt(15.6)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    heading_specs = {
        "Heading 1": (14, WD_ALIGN_PARAGRAPH.CENTER, 0, 6),
        "Heading 2": (12, WD_ALIGN_PARAGRAPH.LEFT, 6, 0),
        "Heading 3": (10.5, WD_ALIGN_PARAGRAPH.LEFT, 6, 0),
    }
    for name, (size, alignment, before, after) in heading_specs.items():
        style = doc.styles[name]
        _set_font(style, "黑体", "Times New Roman", size, bold=True)
        style.paragraph_format.alignment = alignment
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = doc.styles["Caption"]
    _set_font(caption, "宋体", "Times New Roman", 10, bold=True)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.line_spacing = Pt(15)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_together = True


def _add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(_normalise_labels(text), level=level + 1)


def _add_para(doc: Document, text: str, *, indent: bool = False, align: str = "justify") -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(15.6)
    p.add_run(_normalise_labels(text))


def _add_figure(doc: Document, path: Path, caption: str, *, width_cm: float = 14.0) -> None:
    if not path.exists():
        _add_para(doc, f"[图片缺失: {path.name}]", align="center")
    else:
        doc.add_picture(str(path), width=Cm(width_cm))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last.paragraph_format.space_before = Pt(3)
        last.paragraph_format.space_after = Pt(0)
        last.paragraph_format.keep_with_next = True
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(_normalise_labels(caption))


def _add_table_caption(doc: Document, caption: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.keep_with_next = True
    p.add_run(_normalise_labels(caption))


def _add_table(
    doc: Document,
    header: list[str],
    rows: list[list[str]],
    *,
    col_widths_cm: list[float] | None = None,
    font_size_pt: float = 8.5,
    line_spacing_pt: float = 12,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "7F7F7F")
        borders.append(border)
    table._tbl.tblPr.append(borders)
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "DDEBF7")
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = str(val)
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_pr.append(repeat)
    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        row_pr.append(cant_split)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(line_spacing_pt)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(font_size_pt)
                    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    if col_widths_cm:
        for ci, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[ci].width = Cm(w)


def _math_run(text: str) -> OxmlElement:
    run = OxmlElement("m:r")
    value = OxmlElement("m:t")
    value.text = text
    run.append(value)
    return run


def _math_subscript(base: str, subscript: str) -> OxmlElement:
    item = OxmlElement("m:sSub")
    base_node = OxmlElement("m:e")
    base_node.append(_math_run(base))
    sub_node = OxmlElement("m:sub")
    sub_node.append(_math_run(subscript))
    item.extend((base_node, sub_node))
    return item


def _math_superscript(base: str, superscript: str) -> OxmlElement:
    item = OxmlElement("m:sSup")
    base_node = OxmlElement("m:e")
    base_node.append(_math_run(base))
    sup_node = OxmlElement("m:sup")
    sup_node.append(_math_run(superscript))
    item.extend((base_node, sup_node))
    return item


def _math_fraction(numerator: list[OxmlElement], denominator: list[OxmlElement]) -> OxmlElement:
    fraction = OxmlElement("m:f")
    numerator_node = OxmlElement("m:num")
    numerator_node.extend(numerator)
    denominator_node = OxmlElement("m:den")
    denominator_node.extend(denominator)
    fraction.extend((numerator_node, denominator_node))
    return fraction


def _formula_elements(number: str) -> list[OxmlElement]:
    """Return structured OMML for the four numbered equations in this chapter."""
    if number == "(4-1)":
        return [
            _math_subscript("τ", "e"),
            _math_run("(t) = "),
            _math_fraction(
                [_math_subscript("d", "e")],
                [_math_subscript("v", "eff"), _math_run("(risk(t), env(t))")],
            ),
        ]
    if number == "(4-2)":
        parts: list[OxmlElement] = [_math_run("C(path) = ")]
        for index, (weight, term) in enumerate(
            (("d", "d"), ("t", "τ"), ("r", "risk"), ("v", "Δv"), ("u", "u"))
        ):
            if index:
                parts.append(_math_run(" + "))
            parts.extend((_math_subscript("w", weight), _math_run("·∑ ")))
            if term == "d":
                parts.append(_math_subscript("d", "e"))
            elif term == "τ":
                parts.extend((_math_subscript("τ", "e"), _math_run("(t)")))
            elif term == "risk":
                parts.append(_math_run("risk(t)"))
            elif term == "Δv":
                parts.extend((_math_subscript("Δv", "e"), _math_run("(t)")))
            else:
                parts.extend((_math_subscript("u", "e"), _math_run("(t)")))
        return parts
    if number == "(4-3)":
        return [_math_run("R(x, y, "), _math_subscript("t", "arrive"), _math_run(")")]
    if number == "(4-4)":
        return [
            _math_superscript("t", "*"),
            _math_run(" = "),
            _math_subscript("t", "depart"),
            _math_run(" + "),
            _math_subscript("τ", "e"),
            _math_run("("),
            _math_superscript("t", "*"),
            _math_run(")"),
        ]
    raise ValueError(f"unsupported formula number: {number}")


def _formula_fallback_text(number: str) -> str:
    return {
        "(4-1)": "τₑ(t) = dₑ / v_eff(risk(t), env(t))",
        "(4-2)": (
            "C(path) = w_d·∑dₑ + w_t·∑τₑ(t) + w_r·∑risk(t) + "
            "w_v·∑Δvₑ(t) + w_u·∑uₑ(t)"
        ),
        "(4-3)": "R(x, y, t_arrive)",
        "(4-4)": "t* = t_depart + τₑ(t*)",
    }[number]


def _add_formula(doc: Document, _body: str, number: str) -> None:
    """Insert a native Word equation (OMML) with a right-aligned number."""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Cm(0.1), Cm(14.1), Cm(1.2))
    for cell, width in zip(table.rows[0].cells, widths, strict=True):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
        if cell_width is not None:
            cell_width.set(qn("w:w"), str(int(width.twips)))
            cell_width.set(qn("w:type"), "dxa")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    table._tbl.tblPr.append(borders)
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    row_properties.append(cant_split)

    equation_paragraph = table.rows[0].cells[1].paragraphs[0]
    equation_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation_paragraph.paragraph_format.space_before = Pt(3)
    equation_paragraph.paragraph_format.space_after = Pt(3)
    compatibility_namespace = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    alternate = etree.Element(f"{{{compatibility_namespace}}}AlternateContent")
    choice = etree.Element(f"{{{compatibility_namespace}}}Choice")
    choice.set("Requires", "w16cex")
    equation = OxmlElement("m:oMath")
    equation.extend(_formula_elements(number))
    choice.append(equation)
    fallback = etree.Element(f"{{{compatibility_namespace}}}Fallback")
    fallback_run = OxmlElement("w:r")
    fallback_properties = OxmlElement("w:rPr")
    fallback_fonts = OxmlElement("w:rFonts")
    fallback_fonts.set(qn("w:ascii"), "Cambria Math")
    fallback_fonts.set(qn("w:hAnsi"), "Cambria Math")
    fallback_properties.append(fallback_fonts)
    fallback_size = OxmlElement("w:sz")
    fallback_size.set(qn("w:val"), "19")
    fallback_properties.append(fallback_size)
    fallback_run.append(fallback_properties)
    fallback_text = OxmlElement("w:t")
    fallback_text.text = _formula_fallback_text(number)
    fallback_run.append(fallback_text)
    fallback.append(fallback_run)
    alternate.extend((choice, fallback))
    equation_paragraph._p.append(alternate)

    number_paragraph = table.rows[0].cells[2].paragraphs[0]
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_paragraph.paragraph_format.space_before = Pt(3)
    number_paragraph.paragraph_format.space_after = Pt(3)
    number_run = number_paragraph.add_run(number)
    number_run.font.name = "Times New Roman"
    number_run.font.size = Pt(10.5)


def build() -> None:
    if not REFERENCE_DOCX.is_file():
        raise FileNotFoundError(f"reference document not found: {REFERENCE_DOCX}")
    doc = Document(str(REFERENCE_DOCX))
    _clear_body(doc)
    _configure_document(doc)

    _build_title(doc)
    _build_4_1(doc)
    _build_4_2(doc)
    _build_4_3(doc)
    _build_4_4(doc)
    _build_4_5(doc)
    _build_4_6(doc)
    _build_4_7(doc)
    _build_4_8(doc)
    _build_4_9(doc)
    _build_4_10(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"wrote {OUTPUT}")


# ---------------------------------------------------------------------------
# 章标题
# ---------------------------------------------------------------------------
def _build_title(doc: Document) -> None:
    h = doc.add_heading("第4章 风险约束时变航路规划与动态重规划方法", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# 4.1 问题描述与总体技术框架
# ---------------------------------------------------------------------------
def _build_4_1(doc: Document) -> None:
    _add_heading(doc, "4.1 问题描述与总体技术框架", level=1)

    _add_heading(doc, "4.1.1 北极航行动态气象导航的科学问题", level=2)
    _add_para(
        doc,
        "第三章建立了多源环境数据融合、综合风险表征、多时效风险预测与不确定性"
        "输出方法。在此基础上，第四章进一步解决“已知未来不同位置的风险之后，"
        "如何把这些风险真正转化为船舶能够执行的航路和导航建议”这一问题。对"
        "于北极航道而言，风险预测结果只有进入航路代价计算和动态决策过程，才"
        "能真正服务于船舶气象导航。",
        indent=True,
    )
    _add_para(
        doc,
        "本章的核心链条为：环境影响航速 → 风险进入航路代价 → 生成不同航行策略"
        " → 环境变化后主动改线 → 输出多层多目标航路。即系统并不是"
        "在生成风险图后停止，而是将风险场进一步转化为航线规划可以理解的代价"
        "约束，使后续规划模块能够根据不同时间尺度和不同风险水平生成可执行航"
        "行方案。",
        indent=True,
    )
    _add_para(
        doc,
        "由此凝练出本章的科学问题：在时变预测风险约束下，如何构造同时考虑航"
        "程、航行时间、环境风险、航速损失与不确定性的综合航路代价模型，并在"
        "环境变化时生成可执行的动态重规划结果。该问题包含三个子问题：(1) 冰"
        "—浪—风耦合作用下船舶有效航速的机理表达；(2) 时空耦合风险进入航路代"
        "价的建模方法；(3) 在保证解质量与搜索效率的同时，使规划结果可在真实"
        "北极海洋数据上复现且对未知数据 fail-closed。",
        indent=True,
    )

    _add_heading(doc, "4.1.2 总体技术框架与数据流", level=2)
    _add_para(
        doc,
        "本章建立的总体技术框架如图 4-1 所示，由数据输入、边代价建模、航路规"
        "划、验证与输出四个阶段构成"
        "。时变风险场、船舶性能模型与航行任务共同构成规划输入，经时间展开状"
        "态、风险—速度—ETA 耦合边代价与全局失效保护约束建模后，由时间依赖"
        " A* 生成三目标、四层级航路集合；同图基线、消融验证与事件触发动态重"
        "规划共同形成反馈闭环。",
        indent=True,
    )
    _add_figure(doc, FIGURES / "fig-framework.png", "图 4-1 风险约束时变航路规划技术框架")

    _add_para(
        doc,
        "该框架的关键特征在于：风险预测结果并非仅作为背景信息显示，而是进入"
        "航路代价函数（4.2 节）并驱动时空航路优化模型（4.3 节）；规划模块在"
        "时间展开状态图上搜索（4.4 节），并支持事件触发的动态重规划（4.5 节"
        "）；输出保持航路、ETA 与风险摘要的生产者语义，供下游模块验证和展示。",
        indent=True,
    )


# ---------------------------------------------------------------------------
# 4.2 冰—浪—风作用下船舶有效航速与航行代价模型
# ---------------------------------------------------------------------------
def _build_4_2(doc: Document) -> None:
    _add_heading(doc, "4.2 冰—浪—风作用下船舶有效航速与航行代价模型", level=1)

    _add_heading(doc, "4.2.1 环境对有效航速的机理影响", level=2)
    _add_para(
        doc,
        "对于北极航道船舶而言，最短路径并不一定是最优路径。海冰、波浪、风场与"
        "流场共同影响船舶的有效航速、操纵稳定性和航行安全裕度。当某一航段环"
        "境风险较高时，即使几何距离较短，也可能因减速、绕冰、避浪或等待通航"
        "窗口而产生更高的时间代价与安全代价。因此，动态气象导航需要将环境影"
        "响转化为航行代价，而不能仅以距离最短作为判断依据。",
        indent=True,
    )
    _add_para(
        doc,
        "有效航速是船舶在真实环境约束下能够保持的实际航行速度，而不是静水条"
        "件下的理论航速。在开阔水域中，船舶主要受到风浪流影响；在边缘冰区中"
        "，船舶还会受到碎冰、浮冰与冰漂作用影响。当海冰密集度增加时，船舶可"
        "能需要降低航速以减少冰载荷与碰撞风险；当海冰厚度增加时，破冰阻力增"
        "大，推进效率下降；当风浪增强时，纵摇、横摇与偏航风险上升，也可能需"
        "要主动降速；当流场方向与航向不一致时，实际航迹与能耗也会受到影响。",
        indent=True,
    )

    _add_heading(doc, "4.2.2 航段航行时间与有效航速", level=2)
    _add_para(
        doc,
        "航段航行时间不仅取决于航段长度，还取决于环境条件下的有效航速。其基本关系可以表示为：",
        indent=True,
    )
    _add_formula(doc, "τ_e(t) = d_e / v_eff(risk(t), env(t))", "(4-1)")
    _add_para(
        doc,
        "其中，τ_e(t) 表示船舶在时刻 t 出发穿越航段 e 所需的航行时间，d_e 为"
        "航段几何距离，v_eff(·) 为环境约束下的有效航速函数，risk(t) 与 env(t) "
        "分别表示时刻 t 的综合风险值与环境参量。当某一航段风险升高时，即使该"
        "航段距离较短，实际耗时也可能显著增加；反之，绕行距离略长但风险较低"
        "的航段，可能在总航行时间与安全性上更优。",
        indent=True,
    )

    _add_heading(doc, "4.2.3 综合航路代价函数", level=2)
    _add_para(
        doc,
        "航路优化目标不是单纯最短距离，而是综合考虑航程、航行时间、环境风险、航速损失与约束条件后的总代价，其基本形式为：",
        indent=True,
    )
    _add_formula(
        doc,
        "C(path) = w_d · Σ d_e + w_t · Σ τ_e(t) + w_r · Σ risk(t) + w_v · Σ Δv_e(t) + w_u · Σ u_e(t)",  # noqa: E501
        "(4-2)",
    )
    _add_para(
        doc,
        "其中，w_d、w_t、w_r、w_v、w_u 分别为距离代价、时间代价、风险代价、"
        "航速损失代价与不确定性代价的权重；Δv_e(t) 表示环境因素导致的航速损失"
        "；u_e(t) 表示该航段预测可信度的反指标。第三章输出的 risk_mean、"
        "uncertainty、reason_code 与 hard_mask 分别对应式 (4-2) 中的风险代价"
        "输入、保守性调整项、风险来源解释与不可通行区域标记。通过将环境变量"
        "统一转化为航速损失或风险代价，路径搜索结果能够同时考虑航程、风险、"
        "航速损失与约束条件，更符合北极航道气象导航的实际需求。",
        indent=True,
    )


# ---------------------------------------------------------------------------
# 4.3 预测风险约束的时空航路优化模型
# ---------------------------------------------------------------------------
def _build_4_3(doc: Document) -> None:
    _add_heading(doc, "4.3 预测风险约束的时空航路优化模型", level=1)

    _add_heading(doc, "4.3.1 时空耦合风险", level=2)
    _add_para(
        doc,
        "传统静态路径规划通常只考虑某一时刻的空间风险分布，而北极气象导航需"
        "要考虑“船舶何时到达某一位置”。例如，某一区域当前风险较低，但预计 12 "
        "小时后冰缘推进或风浪增强，则船舶抵达该区域时可能已经不适合通行；反"
        "之，某一区域当前风险较高，但未来风险下降，也可能成为后续可选通道。"
        "因此，动态航路优化不能只在当前风险图上寻找低风险路径，而应将空间位"
        "置与到达时间共同纳入计算。",
        indent=True,
    )
    _add_para(doc, "时空耦合风险可表示为：", indent=True)
    _add_formula(doc, "R(x, y, t_arrive)", "(4-3)")
    _add_para(
        doc,
        "其中 x、y 为空间位置，t_arrive 为船舶预计到达该位置的时间。航路规划"
        "模块在评估某一候选航段时，不应只读取当前时刻 R(x, y, t0)，而应读取"
        "船舶到达该区域时的 R(x, y, t_arrive)。这种处理方式使航线规划结果更"
        "接近真实航行过程。",
        indent=True,
    )

    _add_heading(doc, "4.3.2 风险-速度-ETA 三方耦合不动点边代价", level=2)
    _add_para(
        doc,
        "式 (4-1) 与式 (4-3) 表明，边代价 τ_e(t) 不仅是 t 的函数，还通过 risk(t) "
        "与 v_eff 耦合：船舶在时刻 t 出发穿越边 e，其航行时间 τ_e(t) 决定了到"
        "达时刻 t_arrive = t + τ_e(t)，而 t_arrive 又决定了该边在到达时刻的风"
        "险与可达速度，从而再次影响 τ_e。这种三方耦合使边代价求解成为一个不"
        "动点问题：",
        indent=True,
    )
    _add_formula(doc, "t* = t_depart + τ_e(t*)", "(4-4)")
    _add_para(
        doc,
        "即需要找到一个 t*，使得船舶在 t* 出发的航行时间 τ_e(t*) 恰好使其到达"
        "时刻 t_depart + τ_e(t*) 等于 t*。本章采用固定两轮精化逼近该不动点："
        "第一轮以出发时刻的风险场估算 τ_e，第二轮以第一轮到达时刻的风险场修"
        "正 τ_e。固定两轮精化是当前正式求值语义，但它并不构成一般收敛性证明；"
        "本章在受控振荡场中发现不动点可能不存在，因此相关安全改进仅作为默认"
        "关闭的研究方法讨论（详见 4.9.2 节）。",
        indent=True,
    )

    _add_heading(doc, "4.3.3 硬约束条件与临时硬约束", level=2)
    _add_para(
        doc,
        "硬约束是航路规划中必须满足的条件，不能通过降低权重或接受更高代价来"
        "忽略。系统中的硬约束主要包括陆地、浅水区、禁航区、长期限制区以及明"
        "显不可通行的高风险区域。这些区域通过 hard_mask 进行标记，后续航线规"
        "划模块在搜索路径时应直接排除。除静态硬约束外，极端环境条件也可以在"
        "特定情况下转化为临时硬约束：当海冰密集度、冰厚或风浪风险超过船舶"
        "可接受范围时，该区域在对应时效内可被视为不可优先通行区域。",
        indent=True,
    )

    _add_heading(doc, "4.3.4 不同风险偏好下的目标函数权重", level=2)
    _add_para(
        doc,
        "不同航行任务可能对应不同的风险偏好。在安全优先模式下，系统提高风险"
        "代价与不确定性代价权重，优先选择更稳定、更保守的通道；在效率优先模"
        "式下，系统适当增加距离与时间代价权重，在可接受风险范围内缩短航程；"
        "在均衡模式下，系统综合考虑安全、效率与可执行性，生成适合常规航行决"
        "策的推荐路径。本章实验设置三组目标权重，分别对应 fastest（效率优先"
        "）、low_risk（安全优先）与 recommended（均衡推荐）三种航行偏好（详见"
        "表 4-1）。",
        indent=True,
    )


# ---------------------------------------------------------------------------
# 4.4 时间依赖 A* 算法与多模式气象航路生成
# ---------------------------------------------------------------------------
def _build_4_4(doc: Document) -> None:
    _add_heading(doc, "4.4 时间依赖 A* 算法与多模式气象航路生成", level=1)

    _add_heading(doc, "4.4.1 时间展开状态图上的 A* 搜索", level=2)
    _add_para(
        doc,
        "本章以 A* 为确定性搜索骨架，但搜索对象不是静态几何网格，而是由空"
        "间节点、离散到达时间桶和入射航向编码构成的时间展开状态 (node, "
        "time_bucket, heading_code)。时间维度使同一空间节点在不同抵达时刻能保"
        "留不同的风险和航速语义，航向维度则使转向代价能进入状态转移；因而该方"
        "法不是将静态 A* 直接套用于风险底图。每条候选边都调用式(4-1)至式(4-4)"
        "实现风险—速度—ETA 联合求值，再由式(4-2)累积对应目标的综合代价。A* "
        "按 f=g+h 的优先级展开状态，其中 g 为已付出代价，h 为不高估剩余代价的"
        "启发式下界。",
        indent=True,
    )

    _add_heading(doc, "4.4.2 admissible 启发式与最优性保证", level=2)
    _add_para(
        doc,
        "启发式下界仅使用目标间的几何下界、物理可达的最大航速与非负代"
        "价项，因而不高估真实剩余代价。这一设计将“加速”与“修改问题定"
        "义”分离：A* 与无信息 Dijkstra 共享同一时间展开图、边评估器、风险拒"
        "绝语义与终止条件，两者的差别只是节点展开顺序。4.7.1 节的真实 24 小时"
        "实验中，两者总代价满足 cost_identical=True，同时本文 A* 的运行时间为"
        " Dijkstra 的 1/2.31～1/7.84。因此，本章可以在“同一离散状态图”边界内"
        "同时给出代价一致性与搜索效率证据，但不把这一结果外推为连续海洋模型的"
        "全局最优性。",
        indent=True,
    )

    _add_heading(doc, "4.4.3 多模式航路与多时效层级", level=2)
    _add_para(
        doc,
        "系统支持多模式航路（安全优先 / 效率优先 / 均衡推荐）与多时效层级（全"
        "航程 / 24—72 小时主通道 / 0—24 小时滚动 / 0—6 小时可执行段）的组合，"
        "产出 12 条候选航路。全航程结果用于整体参考航线判断，24—72 小时结果"
        "用于主通道可持续性判断，0—24 小时结果用于滚动优化动态航线，0—6 小时"
        "结果用于执行层面的短时修正。这种分层机制使系统既能保持全局航线稳定"
        "，又能在局部风险变化时及时调整。",
        indent=True,
    )

    _add_heading(doc, "4.4.4 算法创新性与设计合理性", level=2)
    _add_para(
        doc,
        "本章的方法贡献不在于重新命名 A* 搜索，而在于将北极航行问题中原本"
        "分散的四类语义统一到一个可复现的时间依赖规划闭环中：一是以 (node, "
        "time_bucket, heading_code) 显式表示“何时、以何航向抵达”；二是在单边"
        "评估中闭合风险、环境降速与 ETA；三是将缺失覆盖、身份不匹配、hard_mask "
        "与超界风险作为 fail-closed 硬拒绝，而不是当作零风险；四是在同一计算"
        "语义下独立生成四个时效层与三类偏好的 12 条航路，使全局参考、滚动决策"
        "与近期执行能够在一致风险尺度上对照。",
        indent=True,
    )
    _add_para(
        doc,
        "上述创新性的核心是“对象建模 + 语义约束 + 可验证搜索”，而非仅用更多"
        "参数换取局部速度。对所有加速候选，本章固定风险输入、船舶模型、网格、"
        "目标函数和失败语义，并同时检查路线一致性、运行时间、P95、命中率、RSS "
        "与真实输入前提。4.9.2 节进一步表明，即使候选在某个局部指标上获得"
        "数十个百分点的改善，只要其正确性前提、单元回归或资源上界未通过，就不改"
        "变当前默认。这种可证伪的候选筛选机制是本方法的合理性与工程竞争力来"
        "源之一。",
        indent=True,
    )


# ---------------------------------------------------------------------------
# 4.5 事件触发动态重规划与规划结果接口
# ---------------------------------------------------------------------------
def _build_4_5(doc: Document) -> None:
    _add_heading(doc, "4.5 事件触发动态重规划与规划结果接口", level=1)

    _add_heading(doc, "4.5.1 事件触发机制与触发阈值", level=2)
    _add_para(
        doc,
        "北极航道的风险不是固定不变的，冰缘推进、冰漂变化、风浪增强、能见度下"
        "降或通航情况变化都可能导致原计划航线风险升高。当系统识别到原航线穿"
        "越高风险区域、低风险通道收窄或不可航行区域扩大时，应触发航路重新评"
        "估。典型触发事件包括：原航线前方海冰密集度明显升高、冰缘向航线方向"
        "推进、冰漂方向对航线形成挤压、风速或浪高超过安全阈值、能见度快速下"
        "降、流场导致偏航风险增加、浅水或限制区约束影响当前航线等。",
        indent=True,
    )
    _add_para(
        doc,
        "动态重规划并不意味着频繁改变航线，而是在风险变化具有实际影响时才触"
        "发。过于频繁的改线会增加航行操作复杂度，也可能导致路径不稳定。因此"
        "，系统设置触发阈值与稳定性约束：当风险升高幅度较小或影响区域较远时"
        "，仅更新风险提示；当风险升高区域与未来航段重合，并可能影响船舶安全"
        "或航速时，才触发重新规划；当高风险区域进入 0—6 小时可执行航线范围时"
        "，应优先生成短时避险建议。",
        indent=True,
    )

    _add_heading(doc, "4.5.2 多层航路结果与关键风险摘要", level=2)
    _add_para(
        doc,
        "规划结果按全航程、主走廊、滚动窗口与可执行段四个层级组织，并在每个"
        "层级给出最快、低风险与均衡推荐三类目标。除航路几何外，结果还保留到"
        "达时间、航段风险、约束状态与重规划标志，形成 12 条航路及其关键风险"
        "摘要。该组织方式既支持长期航路选择，也支持滚动窗口内的局部调整。",
        indent=True,
    )

    _add_heading(doc, "4.5.3 航路规划结果的下游接口", level=2)
    _add_para(
        doc,
        "本章输出由多层多目标航路、到达时间、航段风险与规划状态构成，并以标"
        "准化结果接口供后续展示与导航决策模块使用。风险场属于上游规划输入，"
        "航路可视化与弱通信呈现属于下游消费过程；二者均不在本章内重新计算。"
        "这种边界既保持了规划结果的可追溯性，也避免下游展示改变航路与风险语义。",
        indent=True,
    )


# ---------------------------------------------------------------------------
# 4.6 仿真实验设计
# ---------------------------------------------------------------------------
def _build_4_6(doc: Document) -> None:
    _add_heading(doc, "4.6 仿真实验设计", level=1)

    _add_heading(doc, "4.6.1 实验目的与对比对象", level=2)
    _add_para(
        doc,
        "本章实验目的有三：(1) 验证时间依赖 A* 相对于无信息 Dijkstra 的搜索"
        "效率优势，且不牺牲解质量；(2) 验证时变风险感知规划相对于静态场规划"
        "的航线质量优势；(3) 通过消融研究与反向论证，验证各模块的贡献与当前"
        "实现未被超越的性质。对比对象包括四种算法：本文时间依赖 A*（Proposed"
        "）、无信息 Dijkstra、静态场 A*（Static-field）、风险无关基线（Risk-"
        "blind）。",
        indent=True,
    )

    _add_heading(doc, "4.6.2 数据集与实验变量", level=2)
    _add_para(
        doc,
        "实验数据包括真实北极冬航窗口的 145 帧风险预报序列（holdout 2026-02-22 "
        "与 development 2026-03-22 两个独立窗口，网格 31×11，单一走廊）与四档"
        "合成算例（5×7×7、9×13×13、13×21×25、17×29×37）。真实数据用于质量与"
        "效率对比，合成数据用于可扩展性分析。实验变量为算法（4 档）与目标函"
        "数（fastest / low_risk / recommended 三档）。",
        indent=True,
    )

    _add_heading(doc, "4.6.3 算法公平性与实验统一设置", level=2)
    _add_para(doc, "为使性能比较具有可归因性，四种算法共享以下实验设置，详见表 4-1。", indent=True)
    _add_table_caption(doc, "表 4-1 算法公平性与实验统一设置")
    _add_table(
        doc,
        ["项目", "本文 A*", "Dijkstra", "静态场", "风险无关"],
        [
            ["起终点", "一致", "一致", "一致", "一致"],
            ["网格", "一致", "一致", "一致", "一致"],
            ["船舶性能模型", "一致", "一致", "一致", "一致"],
            ["边评估器", "一致", "一致", "一致", "一致"],
            ["时间离散桶", "一致", "一致", "一致", "一致"],
            ["硬约束(fail-closed)", "是", "是", "是", "是"],
            ["风险信息", "全程使用", "同图无信息", "仅出发时刻帧", "权重置零"],
            ["时变信息", "时间展开图", "同图", "冻结为静态", "风险项取消"],
            ["启发式", "可采纳", "关闭", "可采纳", "可采纳"],
            ["目标函数", "三目标一致", "一致", "一致", "风险与不确定性权重置零"],
            ["硬件", "统一", "统一", "统一", "统一"],
            ["数据集", "同源", "同源", "同源", "同源"],
            ["终止条件", "最大扩展数250000", "同", "同", "同"],
        ],
    )
    _add_para(
        doc,
        "由表 4-1 可知，四种算法仅改变预先定义的单一对比因素，其余输入、约束"
        "与硬件环境保持一致。因此，各组差异可归因于对应的搜索策略、时变信息"
        "或风险项，而非无关实验条件。其中，Risk-blind 是目标函数消融（仅将 risk 与 "
        "uncertainty 权重置零，其余权重保持不变），其搜索效率指标（扩展数、"
        "加速比）不作为优势对比，仅比较风险/时间/航程权衡。",
        indent=True,
    )


# ---------------------------------------------------------------------------
# 4.7 算法性能对比结果与分析
# ---------------------------------------------------------------------------
def _build_4_7(doc: Document) -> None:
    _add_heading(doc, "4.7 算法性能对比结果与分析", level=1)

    _add_heading(doc, "4.7.1 搜索效率对比（vs 无信息 Dijkstra）", level=2)
    _add_para(
        doc,
        "真实 24 小时航段上，本文 A* 与无信息 Dijkstra 的搜索效率对比结果如表 4-2 所示。",
        indent=True,
    )
    _add_table_caption(doc, "表 4-2 搜索效率对比（vs 无信息 Dijkstra，真实 24h）")
    _add_table(
        doc,
        ["算例", "目标", "扩展数(本文)", "扩展数(Dijkstra)", "扩展减少", "加速比", "代价相同"],
        [
            ["holdout", "最快", "653", "4864", "-86.6%", "7.84×", "✓"],
            ["holdout", "低风险", "1828", "4904", "-62.7%", "2.77×", "✓"],
            ["holdout", "推荐", "954", "4845", "-80.3%", "5.17×", "✓"],
            ["development", "最快", "567", "4185", "-86.5%", "7.36×", "✓"],
            ["development", "低风险", "1864", "4211", "-55.7%", "2.31×", "✓"],
            ["development", "推荐", "945", "4173", "-77.4%", "4.65×", "✓"],
        ],
    )
    _add_para(
        doc,
        "由表 4-2 可知，在真实 145 帧风险预报序列上，本文 A* 的节点扩展数较无"
        "信息 Dijkstra 减少 55.7%～86.6%，加速 2.31×～7.84×。更关键的是，全"
        "部 6 个对比单元的总代价严格一致（cost_identical=True），这表明启发"
        "式搜索在保持解质量的前提下有效压缩了搜索空间。其原因在于 admissible "
        "启发式保证不高估真实剩余代价，使 A* 与 Dijkstra 在同一时间展开状态图"
        "上返回相同的最优解，但 A* 通过启发式优先展开更有希望的节点，从而减"
        "少无谓扩展。",
        indent=True,
    )

    _add_heading(doc, "4.7.2 航线质量对比（vs 静态场规划）", level=2)
    _add_para(
        doc,
        "真实 24 小时航段上，本文 A* 与静态场规划（将所有预报帧冻结为出发时刻帧）的航线质量对比结果如表 4-3 所示。",  # noqa: E501
        indent=True,
    )
    _add_table_caption(doc, "表 4-3 航线质量对比（vs 静态场规划，真实 24h）")
    _add_table(
        doc,
        ["算例", "目标", "最大风险(本文)", "最大风险(静态)", "降幅", "平均风险降幅"],
        [
            ["holdout", "最快", "0.09007", "0.12832", "-29.8%", "-15.7%"],
            ["holdout", "低风险", "0.07802", "0.09213", "-15.3%", "-15.9%"],
            ["holdout", "推荐", "0.07802", "0.09213", "-15.3%", "-16.5%"],
            ["development", "三目标", "0.18733", "0.24421", "-23.3%", "-10.7%"],
        ],
    )
    _add_para(
        doc,
        "由表 4-3 可知，本文方法的最大航段风险较静态场规划降低 15.3%～29.8%"
        "，平均航段风险降低 10.7%～16.5%。值得注意的是，在 holdout 窗口的 "
        "fastest 目标下，本文方法与静态场规划的航程完全相同（397.4 km），但"
        "最大风险低 29.8%、平均风险低 15.7%、航行时间还少 0.21 h。这说明优"
        "势并非来自绕远路避险，而是来自在同一条走廊上选出更安全的通行时机——"
        "这正是时变预报与时空耦合风险建模（4.3 节）的价值所在。",
        indent=True,
    )

    _add_heading(doc, "4.7.3 可扩展性分析（合成规模曲线）", level=2)
    _add_para(
        doc,
        "为验证算法随问题规模的可扩展性，在四档合成算例上进行对比，运行时间随网格规模的变化如图 4-2 所示。",  # noqa: E501
        indent=True,
    )
    _add_figure(
        doc, FIGURES / "fig-runtime-scale-log.png", "图 4-2 运行时间随网格规模变化（双对数）"
    )
    _add_para(
        doc,
        "由图 4-2 可知，本文 A*（实线）与无信息 Dijkstra（虚线）的运行时间均"
        "随网格格点数对数线性增长，但本文 A* 始终位于 Dijkstra 下方。在 fastest "
        "目标下，加速比从 5.67×（5×7×7）单调增大至 17.58×（17×29×37），表明"
        "问题规模越大，启发式加速的绝对收益越显著。该结果说明本文方法在更大"
        "规模的真实北极航道网格上仍能保持效率优势。",
        indent=True,
    )

    _add_heading(doc, "4.7.4 效率-质量二维权衡", level=2)
    _add_para(
        doc,
        "为进一步验证“效率提升不以牺牲解质量为代价”，绘制运行时间与总代价、最大风险的二维散点图，分别如图 4-3 与图 4-4 所示。",  # noqa: E501
        indent=True,
    )
    _add_figure(doc, FIGURES / "fig-runtime-cost.png", "图 4-3 运行时间—总代价关系（真实 24h）")
    _add_figure(doc, FIGURES / "fig-runtime-risk.png", "图 4-4 运行时间—最大风险关系（真实 24h）")
    _add_para(
        doc,
        "由图 4-3 可知，本文 A*（蓝色）与 Dijkstra（橙色）在总代价轴上几乎重"
        "合，但在运行时间轴上横向拉开约一个数量级，证明启发式加速未牺牲解代"
        "价。由图 4-4 可知，本文 A* 与 Dijkstra 共享同一目标函数并得到相同总代"
        "价，其风险点相互重合；图中风险差异主要用于比较 Static-field 与 Risk-"
        "blind 的风险变化，而不是声称 A* 相对 Dijkstra 降低风险。需说明的是，development 窗"
        "口上 Risk-blind 与 recommended 路线完全相同（n=1 有效样本，详见 4.9.1 "
        "节），该结果仅用于现象说明，不构成统计意义上的性能结论。",
        indent=True,
    )


# ---------------------------------------------------------------------------
# 4.8 风险分布与逐段风险时序分析
# ---------------------------------------------------------------------------
def _build_4_8(doc: Document) -> None:
    _add_heading(doc, "4.8 风险分布与逐段风险时序分析", level=1)

    _add_heading(doc, "4.8.1 真实 24h 逐段风险时间序列", level=2)
    _add_para(
        doc,
        "为直观展示本文方法对风险峰值的抑制效果，绘制真实 24 小时航段的逐段风险时间序列，如图 4-5 所示。",  # noqa: E501
        indent=True,
    )
    _add_figure(
        doc,
        FIGURES / "fig-risk-timeseries.png",
        "图 4-5 真实 24h 航段逐段风险序列（推荐目标）",
    )
    _add_para(
        doc,
        "由图 4-5 可知，在 holdout 窗口（左子图），本文 A*（蓝色）的逐段风险"
        "始终位于静态场（绿色）与风险无关（红色）曲线下方，且峰值被显著压制"
        "。这表明本文方法不仅在均值上更优，且在极端航段上也能抑制风险峰值。"
        "在 development 窗口（右子图），四条曲线几乎重合，说明该窗口的时变预"
        "报带来的可压低空间较小——这是真实数据特征，非算法缺陷。该结果仅用于"
        "现象说明（n=1），不构成统计意义上的性能结论。",
        indent=True,
    )

    _add_heading(doc, "4.8.2 风险分布箱线分析", level=2)
    _add_para(
        doc,
        "为进一步量化风险分布特征，绘制真实 24 小时航段的风险分布箱线图，如图 4-6 所示。",
        indent=True,
    )
    _add_figure(
        doc,
        FIGURES / "fig-risk-distribution.png",
        "图 4-6 真实 24h 航段风险分布箱线图（推荐目标）",
    )
    _add_para(
        doc,
        "由图 4-6 可知，在 holdout 窗口，本文 A* 与 Dijkstra 的箱子较窄且位"
        "于低风险区间，而静态场与风险无关的箱子更宽且上移，表明后两者不仅在"
        "均值上更差，且风险分布更分散、极端风险更高。这说明本文方法在多数航"
        "段下都更好，且极端风险得到抑制，而不仅是均值意义上的改善。",
        indent=True,
    )


# ---------------------------------------------------------------------------
# 4.9 消融、鲁棒性与反向论证
# ---------------------------------------------------------------------------
def _build_4_9(doc: Document) -> None:
    _add_heading(doc, "4.9 消融、鲁棒性与反向论证", level=1)

    _add_heading(doc, "4.9.1 模块消融研究", level=2)
    _add_para(
        doc,
        "为回答“本文性能提升到底是哪一部分带来的”，进行模块消融研究，结果如表 4-4 所示。",
        indent=True,
    )
    _add_table_caption(doc, "表 4-4 模块消融研究（真实 24h 匹配对照）")
    _add_table(
        doc,
        ["消融档", "描述", "运行时间", "扩展数", "代价/质量", "样本"],
        [
            [
                "完整模型",
                "完整模型",
                "2012.3 ms(dev) / 2147.5 ms(hold)",
                "567 / 653",
                "匹配对照基准",
                "n=2窗口",
            ],
            [
                "去风险项",
                "风险与不确定性权重置零",
                "与Full接近",
                "同图",
                "dev: 0%; hold: 风险+10.4%/+15.4%",
                "n=1有效",
            ],
            [
                "去启发式",
                "Dijkstra",
                "14818.9 ms(dev) / 16839.5 ms(hold)",
                "4185 / 4864",
                "代价严格一致",
                "n=2窗口",
            ],
            [
                "去时变信息",
                "静态场",
                "≈Full",
                "≈Full",
                "hold: 航程同397.4 km; 风险+15.3%/+29.8%",
                "n=2窗口",
            ],
        ],
    )
    _add_para(
        doc,
        "由表 4-4 可知：(1) No risk（风险无关）在 development 窗口与 Full 完全"
        "相同（路线重合），说明 risk 权重在该窗口不起决定性作用；在 holdout "
        "窗口风险高 10.4%/15.4% 而时间仅快 0.5%，证明“使用风险场”带来 15.4% "
        "峰值风险降低而仅多 0.5% 时间。(2) No heuristic（Dijkstra）扩展数增加 "
        "6～7 倍、耗时增加 7～8 倍，但代价严格相同，证明启发式加速不牺牲最优性"
        "。(3) No temporal（静态场）在 holdout 上航程与 Full 完全相同但风险高 "
        "15.3%/29.8%，证明“使用时变预报”的价值在于在同一条走廊上选出更安全的"
        "通行时机。需说明的是，development 窗口上 Risk-blind 与 recommended 路"
        "线完全相同，故有效样本 n=1（holdout），该结果仅用于现象说明，不构成"
        "统计意义上的性能结论。",
        indent=True,
    )
    _add_para(
        doc,
        "此外，关于“no_replanning / no_correction”独立消融档：正式规划路径默认"
        "就是单次规划，重规划基线是单独验证产物，与默认路径正交，在正式控制下"
        "不存在“再剥一层重规划”的独立语义，故本章消融研究严格只列上述 4 档"
        "。",
        indent=True,
    )

    _add_heading(doc, "4.9.2 改进候选门禁反向论证", level=2)
    _add_para(
        doc,
        "本章对 6 个改进候选（FIFO 支配剪枝、轨迹复用、SMO-A* 共享记忆化、"
        "ARA* anytime 备选、bounded LRU 风险采样缓存、non-FIFO exact-arrival 与"
        "完整 Pareto frontier）进行了完整评估。结果表明，没有一个候选通过最"
        "终门禁、没有一个被启用，如图 4-7 所示。",
        indent=True,
    )
    _add_figure(doc, FIGURES / "fig-funnel.png", "图 4-7 改进候选门禁漏斗：当前实现未被超越")
    _add_table_caption(doc, "表 4-5 改进候选的正向观测、失败门禁与当前结论")
    _add_table(
        doc,
        ["候选方案", "已观测的正向证据", "未通过的门禁与实证数据", "当前结论"],
        [
            [
                "P0.1 FIFO 支配剪枝",
                "synthetic small/medium 各60个证书化剪枝；compute_ms改善58.22%/76.52%",
                "两套145帧真实输入发现43500/40776个interval级FIFO违反，前提性门禁失败",
                "真实FIFO支配禁用（REAL_INPUT_FIFO_VIOLATED）",
            ],
            [
                "P0.2 non-FIFO exact-arrival/Pareto",
                "synthetic 90/90通过；18个partial cases、12次transition预剪枝；"
                "真实24h两窗6/6语义/frontier等价",
                "development/holdout真实transition检查271120/499536次，新增预剪枝均0；cgroup强资源上界证据不完整",
                "完成至M34并冻结；不构成默认启用资格",
            ],
            [
                "LTCR-TDA* / P2.1轨迹复用",
                "Winter正式4次重复观测总耗时中位改善47.86%，路线完整性48/48",
                "rolling_0_24h×fastest单元回归5.94%>5%冻结上限；后续短复测样本不足，不能改写原门禁",
                "测量不确定；正式M2失败不变",
            ],
            [
                "SMO-A* 共享记忆化",
                "synthetic small/medium改善55.02%/46.27%；真实路线一致性与P95通过",
                "holdout改善11.22%<15%、命中率14.27%<50%、RSS比3.367>1.10；development命中率19.19%、RSS比3.380",
                "DEFERRED/RETIRED；不进入正式M2",
            ],
            [
                "ARA* anytime备选",
                "medium三目标首解改善44.57%/82.19%/63.16%，gap=0%",
                "small的fastest/recommended首解仅改善4.14%/4.19%<每目标20%，low_risk虽为40.42%仍触发全目标门禁",
                "M0_FAIL/DEFERRED；不进入Winter",
            ],
            [
                "bounded LRU风险采样缓存",
                "direct medium 76.281s→65.012s，中位改善14.77%",
                "RSS增劦约38.6MiB，且未覆盖正式ingress、四层三目标的12航路门禁",
                "EXPERIMENTAL；默认关闭",
            ],
        ],
        col_widths_cm=[2.4, 4.0, 5.4, 2.5],
        font_size_pt=7.5,
        line_spacing_pt=10.5,
    )
    _add_para(
        doc,
        "图4-7与表4-5共同给出的不是简单的“失败清单”，而是一组同源、可证"
        "伪的对抗性消融实验。6个候选分别攻击当前方法的标签数量、重复搜索、风险"
        "采样、首解速度和非FIFO语义处理等瓶颈；其中多个候选曾在synthetic或局部工程"
        "指标上获得40%～80%量级的改善，但仍因真实输入前提、单元回归、跨规模"
        "一致性、缓存命中率或资源上界失败而被拒绝。这说明候选未被启用并非因为"
        "没有尝试改进，而是因为本章要求“快”必须与“语义不变、最差单元不退化、"
        "资源可控”同时成立。",
        indent=True,
    )
    _add_para(
        doc,
        "从反向论证看，P0.1证明真实海洋输入不允许直接套用教科书式FIFO支配；"
        "P0.2在有限真实24h状态域内得到与完整Pareto frontier一致的业务语义，但"
        "没有带来新的transition剪枝，说明当前基线在已验证有限域内没有被该更完整"
        "标签语义改写；P2.1则证明即使总体加速接近48%，也必须因单个业务单元"
        "5.94%回归而停止；SMO-A*与bounded LRU共同说明缓存收益必须与记忆体上界"
        "一起评价；ARA*则说明中等规模的显著改善不能掩盖小规模上两个目标未达"
        "门禁。因此，当前时间依赖A*的竞争力更准确地体现为：在同一风险—速度"
        "—ETA状态图上，它是目前唯一同时保持语义完整、失败关闭、资源可预期与"
        "12航路输出一致性的默认实现。",
        indent=True,
    )
    _add_para(
        doc,
        "需特别说明：该漏斗图与表4-5仅支持“当前默认未被通过全部冻结门禁的"
        "候选超越”，不支持“已证明为所有传统算法中的性能最优解”。表4-5中的正向"
        "数据是研究贡献和候选设计有效性的证据，失败门禁则是不将局部收益误报为"
        "已实现算法优势的证据。",
        indent=True,
    )

    _add_heading(doc, "4.9.3 真实数据违反 FIFO 的科学发现", level=2)
    _add_para(
        doc,
        "教科书时依赖最短路算法（time-dependent A* / Dijkstra）的时间支配剪"
        "枝成立的前提是 FIFO 性质：出发越晚，到达不更早。本章用分区证据扫描"
        "真实北极冬航 145 帧输入，发现该前提在真实海洋数据上不成立，如表 4-6 "
        "所示。",
        indent=True,
    )
    _add_table_caption(doc, "表 4-6 真实海洋数据违反 FIFO 的实证发现")
    _add_table(
        doc,
        ["输入", "有向边×目标", "interval评估", "FIFO违反", "certified probes", "首个反例"],
        [
            [
                "holdout",
                "1388×3",
                "104100",
                "43500(每目标14500)",
                "101958",
                "edge[(0,0),(0,1)], 2026-02-22, 左2.8943h, 右2.7839h",
            ],
            [
                "development",
                "1540×3",
                "115500",
                "40776(每目标13592)",
                "113685",
                "edge[(0,0),(0,1)], 2026-03-22, 左2.9089h, 右2.9064h",
            ],
        ],
    )
    _add_para(
        doc,
        "由表 4-6 可知，跨两独立窗口、每目标约 1.4 万个 interval 级 FIFO 违反"
        "。即“同一时刻出发的左 image 到达时间”比“晚 2 小时出发的右 image”更"
        "晚——出发晚反而到得早。这一发现使教科书时依赖最短路的支配剪枝在真实"
        "数据上不安全，也是本章采取“先证明、后启用”保守工程纪律的直接原因："
        "任何想使用FIFO支配剪枝的候选都必须在真实数据上先证明其前提，而当前真"
        "实输入已拒绝该前提。需说明的是，这是科学发现，不是已解决的问题；正式控制仍采用"
        "固定两轮 ETA 精化（详见 4.3.2 节）。",
        indent=True,
    )

    _add_heading(doc, "4.9.4 诚实性边界与本章结论适用范围", level=2)
    _add_para(
        doc,
        "为避免论文结论被误读，本章明确以下边界：(1) 真实样本仅 2 个独立窗口"
        "，development 窗口上三目标收敛到同一条路线，故质量结论的实际有效样本"
        "为 n=1（holdout）；(2) 效率优势源自 admissible 启发式的正确工程实现，"
        "而非新的搜索算法；(3) “最优性不变”限于同一离散时间展开图，连续海洋模"
        "型上的全局最优性未证明；(4) 静态场基线是构造基线，非某一公开算法的"
        "复现，优势应表述为“使用时变预报”这一设计选择的收益；(5) 扩展数减少不"
        "等于端到端吞吐提升，本章未测完整 12 路线四层规划；(6) 反向论证仅指"
        "“未被超越 + 正确性保守”，不是“性能最优”或“生产级稳定优势”的声明；(7) "
        "本章未画风险场空间热力底图（数据未暴露投影矩阵）、路径地理空间叠加"
        "（网格为抽象索引无投影）、参数敏感性（无系统化扫参数据），相关缺失"
        "已在《第四章视觉证据计划》§4 声明。",
        indent=True,
    )


# ---------------------------------------------------------------------------
# 4.10 本章小结
# ---------------------------------------------------------------------------
def _build_4_10(doc: Document) -> None:
    _add_heading(doc, "4.10 本章小结", level=1)
    _add_para(
        doc,
        "本章针对北极航道动态气象导航的科学问题，建立了风险-速度-ETA 三方耦"
        "合不动点边代价模型与预测风险约束的时空航路优化模型，采用时间依赖 A* "
        "算法在时间展开状态图上搜索多模式气象航路，并支持事件触发的动态重规"
        "划及标准化结果输出。实验结果表明：在真实北极冬航 145 帧风险预报序列"
        "上，相比静态场规划，本文方法的最大航段风险降低 15.3%～29.8%、平均航"
        "段风险降低 10.7%～16.5%，且在 fastest 目标下航程与基线完全相同（397.4 "
        "km）而风险显著更低，说明优势来自对通行时机的选择而非绕行；相比无信息 "
        "Dijkstra，节点扩展数减少 55.7%～86.6%、加速 2.31×～7.84×，且总代价"
        "严格相同，证明启发式加速未牺牲最优性；在合成规模曲线上，加速比随网格"
        "规模单调增大（5.67×→17.58×）。本章还发现真实北极海洋数据违反 FIFO "
        "性质（两窗口 43 500 / 40 776 次 interval 级违反），使教科书时依赖最"
        "短路的支配剪枝在真实数据上不安全，据此采取“先证明、后启用”的保守工程"
        "纪律。这些结果为下一章（第五章）的多模式航路可视化与决策展示提供了"
        "可执行的多模式航路候选、事件触发重规划建议与关键风险摘要接口。",
        indent=True,
    )


if __name__ == "__main__":
    build()
